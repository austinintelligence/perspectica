from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


DOCUMENT_PATH = Path("output/doc/Perspectica - Final Product Design Document.docx")


def paragraph_with_prefix(document: Document, prefix: str) -> Paragraph:
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph beginning with {prefix!r}; found {len(matches)}")
    return matches[0]


def insert_after(anchor: Paragraph, text: str, style: str = "Normal") -> Paragraph:
    new_element = deepcopy(anchor._p)
    for child in list(new_element):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            new_element.remove(child)
    anchor._p.addnext(new_element)
    paragraph = Paragraph(new_element, anchor._parent)
    paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def replace_paragraph(document: Document, prefix: str, text: str) -> None:
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.startswith(prefix)]
    if len(matches) == 1:
        matches[0].text = text
        return
    if len(matches) == 0 and any(paragraph.text == text for paragraph in document.paragraphs):
        return
    raise RuntimeError(f"Expected one paragraph beginning with {prefix!r}; found {len(matches)}")


def main() -> None:
    document = Document(DOCUMENT_PATH)

    replace_paragraph(
        document,
        "Perspectica opens in the Chrome side panel",
        "Perspectica opens in the Chrome side panel. On first use, the user connects ChatGPT "
        "through a clear consent and one-time-code sign-in flow. After the saved session is "
        "ready, Perspectica begins analyzing the active article. The article header and every "
        "section heading appear immediately. Completed results fade into place while the user "
        "continues reading.",
    )

    prompt_version = paragraph_with_prefix(
        document, "Prompt version, model version, extraction version"
    )
    if not any(
        paragraph.text.startswith("Authenticated ChatGPT session")
        for paragraph in document.paragraphs
    ):
        insert_after(
            prompt_version,
            "Authenticated ChatGPT session and the Codex models available to the connected "
            "account for live analysis",
            "List Bullet",
        )

    modular_monolith = paragraph_with_prefix(document, "The system is a modular monolith")
    auth_text = (
        "The live POC uses a consent-first ChatGPT device flow. The extension keeps only a "
        "signed, HttpOnly session cookie. The backend encrypts refreshable session data "
        "before saving it to SQLite, discovers the Codex models available to the connected "
        "account, and sends AI SDK requests through a request-bound proxy. Access and refresh "
        "tokens stay inside the backend authentication handler."
    )
    existing_auth = next(
        (
            paragraph
            for paragraph in document.paragraphs
            if paragraph.text.startswith("The live POC uses a consent-first ChatGPT")
        ),
        None,
    )
    if existing_auth:
        existing_auth.text = auth_text
    else:
        insert_after(
            modular_monolith,
            auth_text,
        )

    stack = document.tables[3]
    for row in stack.rows:
        if row.cells[0].text.strip() == "AI":
            row.cells[1].text = "Vercel AI SDK 7 with Zod"
            row.cells[2].text = (
                "Structured Article Lens, bounded tools, typed section streaming."
            )
        if row.cells[0].text.strip() == "Storage":
            row.cells[2].text = (
                "Encrypted login sessions, article fingerprints, final reports, prompts, "
                "models, metrics, and evaluation runs."
            )

    if not any(row.cells[0].text.strip() == "Authentication" for row in stack.rows):
        backend_row = next(row for row in stack.rows if row.cells[0].text.strip() == "Backend")
        auth_row = stack.add_row()
        auth_row.cells[0].text = "Authentication"
        auth_row.cells[1].text = "Login with ChatGPT device flow"
        auth_row.cells[2].text = (
            "Consent, encrypted server sessions, model discovery, proxy access, and logout."
        )
        stack._tbl.remove(auth_row._tr)
        backend_row._tr.addnext(auth_row._tr)

    replace_paragraph(
        document,
        "packages/article-lens",
        "packages/analysis - Article Lens, claim extraction, pipeline orchestration, and "
        "provider interfaces",
    )
    replace_paragraph(
        document,
        "packages/research",
        "packages/storage - encrypted ChatGPT sessions and the SQLite repositories used by "
        "the backend",
    )
    replace_paragraph(
        document,
        "packages/evaluation",
        "tests and fixtures - expert labels, prompt variants, stability checks, cost, and "
        "latency reports",
    )

    replace_paragraph(
        document,
        "The extension sends one POST request",
        "The extension sends one credentialed POST request to /api/analyze with the extracted "
        "article and a content fingerprint. In live mode, the route verifies the saved ChatGPT "
        "session, discovers the connected account's available Codex models, and selects the "
        "configured or first available model. The route returns a newline-delimited stream of "
        "typed analysis events. The Source List and metadata are emitted first. Compass and "
        "Bias are emitted after the Article Lens passes deterministic validation. Journalist "
        "Context and external evidence are emitted as their source-backed sections finish. "
        "The completed report is stored against the fingerprint so reopening the same article "
        "can load the final result quickly.",
    )

    storage_module = paragraph_with_prefix(document, "Storage module:")
    storage_module.text = (
        "Storage module: Saves encrypted ChatGPT session envelopes with expiration, then saves "
        "the input fingerprint, final report, prompt and model versions, timings, token use, "
        "search counts, and evaluation metadata."
    )
    storage_module.style = "List Bullet"

    pipeline_intro = paragraph_with_prefix(document, "The system uses a structured two-part")
    if not any(
        paragraph.text.startswith("For live analysis, the backend first resolves")
        for paragraph in document.paragraphs
    ):
        insert_after(
            pipeline_intro,
            "For live analysis, the backend first resolves the encrypted ChatGPT session and "
            "selects a model returned by the connected account. AI SDK then calls that model "
            "through the server-side proxy, keeping token material inside the authentication "
            "handler.",
        )

    document.save(DOCUMENT_PATH)


if __name__ == "__main__":
    main()
